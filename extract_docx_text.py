
import docx
import sys

def extract_text(docx_path):
    try:
        doc = docx.Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append(" | ".join(row_text))
        
        return '\n'.join(full_text)
    except Exception as e:
        return str(e)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python extract_docx_text.py <docx_path>")
        sys.exit(1)
    
    path = sys.argv[1]
    print(extract_text(path))
